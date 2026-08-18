from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

ROOT=Path('/mnt/data/aaowasi369v17')
out=ROOT/'AAO_v17_Code_Customization_Guide.docx'

doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(.62); sec.bottom_margin=Inches(.62); sec.left_margin=Inches(.68); sec.right_margin=Inches(.68)

# Theme colors
INK='14231E'; MUTED='5F6F68'; GREEN='0B9B73'; PALE='EAF7F2'; LINE='D8E5DF'; DARK='07110D'

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def borders(cell,color=LINE,size='6'):
    tcPr=cell._tc.get_or_add_tcPr(); b=tcPr.first_child_found_in('w:tcBorders')
    if b is None:
        b=OxmlElement('w:tcBorders'); tcPr.append(b)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el=OxmlElement('w:'+edge); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),size); el.set(qn('w:color'),color); b.append(el)

def set_cell_text(cell,text,bold=False,color=INK,size=8.5):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(text); r.bold=bold; r.font.name='Aptos'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.autofit=False
    for i,h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i],h,True,'FFFFFF',8); shade(t.rows[0].cells[i],GREEN); borders(t.rows[0].cells[i],GREEN)
    trPr=t.rows[0]._tr.get_or_add_trPr(); hdr=OxmlElement('w:tblHeader'); hdr.set(qn('w:val'),'true'); trPr.append(hdr)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            set_cell_text(cells[i],str(val),False,INK,7.9); borders(cells[i]);
            if ri%2: shade(cells[i],'F6FAF8')
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1)
    return t

def heading(text,level=1):
    p=doc.add_paragraph(); p.style=f'Heading {level}'; p.paragraph_format.keep_with_next=True
    r=p.add_run(text); return p

def para(text,bold_prefix=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5); p.paragraph_format.line_spacing=1.08
    if bold_prefix and text.startswith(bold_prefix):
        r=p.add_run(bold_prefix); r.bold=True; p.add_run(text[len(bold_prefix):])
    else:p.add_run(text)
    return p

def bullet(text):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(2); p.add_run(text); return p

def code(text):
    p=doc.add_paragraph(); p.style='Code'; p.paragraph_format.space_after=Pt(4); p.add_run(text); return p

def picture(path,width,caption):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
    p.add_run().add_picture(str(path),width=Inches(width))
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after=Pt(6)
    r=c.add_run(caption); r.italic=True; r.font.size=Pt(7.5); r.font.color.rgb=RGBColor.from_string(MUTED)

# Styles
styles=doc.styles
styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(9.2); styles['Normal'].font.color.rgb=RGBColor.from_string(INK)
for n,size in [('Title',31),('Heading 1',20),('Heading 2',14),('Heading 3',11)]:
    st=styles[n]; st.font.name='Georgia'; st.font.color.rgb=RGBColor.from_string(DARK); st.font.size=Pt(size); st.font.bold=True
    st.paragraph_format.space_before=Pt(8 if n!='Title' else 0); st.paragraph_format.space_after=Pt(5)
if 'Code' not in styles:
    st=styles.add_style('Code',WD_STYLE_TYPE.PARAGRAPH)
else: st=styles['Code']
st.font.name='Liberation Mono'; st.font.size=Pt(7.5); st.font.color.rgb=RGBColor.from_string('14352B')
st.paragraph_format.left_indent=Inches(.18); st.paragraph_format.right_indent=Inches(.08); st.paragraph_format.space_after=Pt(3)

# Header/footer
for section in doc.sections:
    hp=section.header.paragraphs[0]; hp.text='AAO PORTFOLIO · v17 CODE CUSTOMIZATION GUIDE'; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.size=Pt(7); hp.runs[0].font.bold=True; hp.runs[0].font.color.rgb=RGBColor.from_string(GREEN)
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

# Title
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_before=Pt(46); p.paragraph_format.space_after=Pt(10)
r=p.add_run('AAO Portfolio v17'); r.font.name='Georgia'; r.font.size=Pt(35); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(DARK)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(15)
r=p.add_run('Code Customization, Responsiveness & Deployment Guide'); r.font.name='Georgia'; r.font.size=Pt(21); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(GREEN)
para('For Md. Abdullah Al Owasi · Technology Risk · GRC · Security Compliance · TPRM · AI Governance')

box=doc.add_table(rows=4,cols=2); box.autofit=False
for row in box.rows:
    for c in row.cells: borders(c); c.width=Inches(3.3)
entries=[('Release','17.0.0'),('Runtime','Static HTML + CSS + browser JavaScript'),('Production assets','assets/site-v17.css + assets/app-v17.js'),('Build output','out/ — Cloudflare Pages / Vercel compatible static files')]
for i,(a,b) in enumerate(entries):
    set_cell_text(box.cell(i,0),a,True,GREEN,8.5); set_cell_text(box.cell(i,1),b,False,INK,8.5)
    if i%2: shade(box.cell(i,0),'F3F9F6'); shade(box.cell(i,1),'F3F9F6')

para('Line-number rule: the line references in this guide are exact for the delivered v17 source. Any later edit can shift line numbers. The stable way to relocate a section is to search for the named V17-CUSTOMIZE comment, CSS selector, data-* hook, or JavaScript function shown beside the line number.')

heading('1. What controls what',1)
add_table(['Layer','File','Responsibility'],[
('Content / structure','index.html','Text, cards, sections, tables, data hooks, source links, section order.'),
('Visual system','assets/site-v17.css','Typography, themes, layout, container queries, transitions, card geometry, matrix/table modes.'),
('Behavior','assets/app-v17.js','No-click state switching, touch/keyboard fallbacks, deck movement, Dynamic Island, live CISA data.'),
('Release gate','scripts/preflight.mjs + verify-output.mjs','Checks required files, broken references, component counts, forbidden runtimes and production signatures.'),
('Deployment','vercel.json + DEPLOYMENT.md','Static output contract: npm run build → out/.')],[1.35,2.4,3.15])

heading('2. HTML section map — exact edit points',1)
add_table(['Area','Final v17 lines','Stable search marker','What to edit'],[
('Dynamic Island','28–105','data-island','Destinations, labels and order.'),
('Hero + Governance','106–554','V17-CUSTOMIZE HERO','Positioning copy, seven governance stage cards, proof, metrics, signature model.'),
('Governance stage','232–348','data-governance-stage','Card labels and 01–07 controls.'),
('Governance proof','349–430','data-governance-proof','Input / Inspectable proof / Decision state fields.'),
('Signature model','431–554','signature-model-v17','Requirement → Control → Evidence → Exception → Residual risk → Decision.'),
('Executive Value','555–1271','V17-CUSTOMIZE EXECUTIVE VALUE','Value cards and quantitative risk visuals.'),
('Architecture','1272–2037','V17-CUSTOMIZE ARCHITECTURE','Three modules, evidence tables, AI matrix and module decision lattices.'),
('Selected Systems','2038–2229','V17-CUSTOMIZE SELECTED SYSTEMS','Category buttons and ten system cards.'),
('Capabilities','2230–2561','V17-CUSTOMIZE CAPABILITIES','Category rail and capability cards.'),
('Framework Depth','2562–2669','V17-CUSTOMIZE FRAMEWORK DEPTH','Source-linked framework cards.'),
('Live Risk Pulse','2670–2786','V17-CUSTOMIZE LIVE RISK','CISA metrics placeholders, live cards, governance-action path.'),
('Direct Conversation','2787–2890','V17-CUSTOMIZE CONTACT','Recruiting positioning, CTAs and availability.'),
('Footer','2891–3038','V17-CUSTOMIZE FOOTER','Navigation, evidence links and centered Back to Top.')],[1.25,1.0,2.15,2.55])

heading('3. Governance Portfolio — the state machine',1)
para('The first screenshot problem was caused by too many textual layers competing inside the same moving visual. v17 separates the moving stage from the proof area. Only the focused governance card exposes its content; background cards keep depth without readable text.')
picture(ROOT/'validation/v17_final_governance_390.png',3.55,'Final narrow-container Governance Portfolio: moving stage, proof cards, metrics and the 3×2 signature model remain separate.')
add_table(['Change','HTML','CSS','JavaScript'],[
('Stage labels','232–348','governance-stage ~3475+','TRACE + selectStage() ~154–290'),
('Proof text','349–430','governance-proof-v17 ~3511–3555','renderProof() ~184–200'),
('Signature model','431–554','signature-model-v17 ~3558–3610','signature marker sync ~178–231'),
('Animation duration','—','--motion-v17 line ~3411','WAAPI proof transition ~190')],[1.45,1.55,2.15,2.15])
code('const TRACE = [ { stage: "Requirement", input: "…", evidence: "…", decision: "…" }, … ];')
para('Safe edit: change stage substance in TRACE and the matching visible card labels together. Do not add an autoplay timer; all seven states intentionally remain user-controlled.')

doc.add_page_break()
heading('4. Responsive architecture tables — no overlap',1)
para('At wide widths the tables stay real column tables. When a table wrapper itself becomes narrow, the CSS changes each row into a labeled record card. This is based on container width, so it also works in split-screen or narrow embedded layouts on a large monitor.')
picture(ROOT/'validation/v17_ai_table_390.png',2.25,'AI evidence table in its narrow record-card mode; every value has an explicit field label and no horizontal text collision.')
add_table(['Control','Location','How it works'],[
('Table markup','Architecture tables at ~1378, ~1533, ~1688','Each td carries data-label="…".'),
('Container declaration','CSS ~3629–3657','.table-wrap becomes an evidence-table container.'),
('Record-card conversion','CSS ~3658–3699','thead hides; rows become cards; td::before uses data-label.'),
('Overflow sign','JS syncTableHints() ~554–564','The ↔ sign is shown only when real horizontal overflow exists.')],[1.55,2.05,3.6])
para('If you add a new column, add the <th> and the corresponding data-label on every <td>. Missing data-label values will make the narrow mode ambiguous.')

doc.add_page_break()
heading('5. AI risk matrix — 2D plot when space exists, score cards when it does not',1)
picture(ROOT/'validation/doc_ai_matrix_390.png',2.05,'Narrow AI risk view: Risk /15 and Oversight /5 remain explicit while absolute-positioned nodes become clean vertical cards.')
add_table(['Control','Lines / selector','Purpose'],[
('Desktop values','index.html Architecture AI panel','Every node has data-risk and data-oversight values.'),
('Container','CSS ~3705–3711 .heatmap-panel','Measures the matrix panel, not only the browser viewport.'),
('Narrow conversion','CSS @container <=760px ~3712–3741','Removes absolute plot geometry and stacks score cards.'),
('Very narrow','CSS @container <=420px ~3742–3744','Node score moves below label for extra wrapping space.'),
('Pointer inspection','app-v17.js Risk model/AI matrix ~490–549','Nearest node can update on pointer movement; focus/click remain available.')],[1.45,2.45,3.3])

doc.add_page_break()
heading('6. Typography, cards and Direct Conversation',1)
add_table(['Element','CSS edit point','Recommended change'],[
('Body / controls','body.v17 + --ui-v17 ~3409–3465','Change the local UI font stack here; preserve fallbacks and no remote dependency if portability is priority.'),
('Editorial titles','existing display/serif variables + title rules','Keep the serif display treatment for visual authority; change line-height before pushing font-size higher.'),
('System/capability/framework cards','~3748–3769','Shared min-height, flex column and bottom-aligned metadata keep all three deck types visually consistent.'),
('Direct Conversation','~3803–3811 + final rule ~3903–3913','Final max is intentionally below hero scale; edit the last rule only.'),
('Button/icon alignment','island-core + shared controls ~3776–3798','Controls use grid/place-items:center; keep icon boxes square to preserve optical centering.')],[1.6,2.25,3.35])
picture(ROOT/'validation/v17_contact_390.png',2.55,'Direct Conversation at 390px after the final scale correction; the headline no longer dominates the entire viewport.')

doc.add_page_break()
heading('7. Architecture, Systems and Capabilities interactions',1)
para('Desktop pointer movement is an enhancement: moving across the module/category controls switches content without requiring a click. Touch rails synchronize to the nearest/end item after scrolling. Keyboard focus and click remain valid fallbacks.')
add_table(['Interaction','HTML hook','JavaScript'],[
('Architecture modules','data-auto-rail="architecture" ~1286','switchArchitecture() ~354 + adaptive rail setup ~394–422'),
('Selected Systems','data-auto-rail="systems" ~2051','switchProjectFilter() + shared adaptive rail ~366–423'),
('Capabilities','data-auto-rail="capabilities" ~2243','switchSkillFilter() + shared adaptive rail ~367–423'),
('Finite card decks','data-deck-track ~2082 / 2274 / 2592','Finite-deck engine ~427–489; no clones, wrap-around or autoplay.')],[1.55,2.45,3.25])
picture(ROOT/'validation/guide_architecture_crop.png',6.55,'Desktop Architecture: module rail, decision lattice, evidence table and AI risk view stay inside one inspectable workspace.')

doc.add_page_break()
heading('8. Live Public Risk Pulse — current evidence without inventing employer problems',1)
para('The live section uses CISA Known Exploited Vulnerabilities as a public market/risk signal. CISA provides KEV in JSON/CSV form, and its official cisagov/kev-data GitHub repository exists to make the same dataset easier to consume programmatically. CISA states that the mirror is normally synchronized with the canonical source within minutes.')
picture(ROOT/'validation/guide_live_crop.png',1.82,'Live Risk Pulse layout with a deterministic QA feed. Production fetches CISA data on page load and preserves the same layout.')
add_table(['Item','Code location','Behavior'],[
('Section markup','index.html ~2670–2786','Metrics, live cards, source link and four-step governance response path.'),
('Styling','site-v17.css ~3814–3897','Responsive metrics/cards, live/error states, strict light/dark surfaces.'),
('Feed loader','app-v17.js ~594–728','Official CISA GitHub mirror first; canonical CISA JSON second; 5.2s timeout per source.'),
('Calculations','app-v17.js render path ~630–683','Catalog count, additions last 30 days, distinct recent vendors, latest three records.'),
('Failure mode','app-v17.js ~687–701','Shows unavailable state and em dashes; never fabricates current values.')],[1.45,2.3,3.5])
para('Hiring-use rule: public data can show why disciplined Technology Risk/GRC/TPRM work matters, but the site should not say a specific prospective employer suffered a breach unless a reliable public source establishes that fact. Keep the translation at the level of exposure review, ownership, remediation evidence, exception handling and residual-risk decisions.')

doc.add_page_break()
heading('9. Dynamic Island, navigation and scroll safety',1)
add_table(['Control','Lines','Reason'],[
('Island HTML','index.html 28–105','Eight in-page destinations; compact closed state.'),
('Open/close + section state','app-v17.js ~94–147','Tracks current section; Escape/outside click closes it.'),
('Visibility threshold','app-v17.js ~136–142','Appears after the early hero rather than sitting on top of hero content.'),
('Footer suppression','app-v17.js ~135–142','Closes/hides near footer to avoid competing with Back to Top.'),
('Icon centering','site-v17.css ~3776–3795','26×26 grid box; plus/cross center deviation measured at 0px in QA.'),
('Back to Top','index.html ~3020 + JS ~738','Only the footer-centered button scrolls the document to the top.')],[1.7,1.8,3.75])
para('Scroll-jump rule: component state changes do not call scrollIntoView(). Horizontal deck/rail movement uses element.scrollTo() on the component itself. This prevents state changes from dragging the document vertically.')

doc.add_page_break()
heading('10. Deployment and release gate',1)
add_table(['Platform','Configuration','Why it is portable'],[
('Cloudflare Pages','Framework None · npm run build · output out/','Cloudflare Pages can deploy custom/static build output directories.'),
('Vercel','vercel.json: framework null · buildCommand npm run build · outputDirectory out','Vercel supports outputDirectory configuration for the deployment output.'),
('Other static host','Publish out/','Production contains ordinary HTML, CSS, JavaScript and assets; no server runtime is required.')],[1.55,2.85,2.85])
code('npm run check\nnpm run build\nnpm run postbuild')
para('Preflight checks required assets, local references, duplicate IDs, component counts, live-risk signatures, forbidden runtime packages, no recurring setInterval, and no scrollIntoView in production component logic. Postbuild re-checks the generated out/ directory.')

doc.add_page_break()
heading('11. Validation matrix and what “adaptive” means',1)
add_table(['Viewport width','Horizontal page overflow','Known UI overlaps','Unexpected scroll jump','No-click filters'],[
('1440','0 px','0','0 px','Pass'),('1024','0 px','0','0 px','Pass'),('768','0 px','0','0 px','Pass'),('620','0 px','0','0 px','Pass'),('540','0 px','0','0 px','Pass'),('430','0 px','0','0 px','Pass'),('390','0 px','0','0 px','Pass'),('360','0 px','0','0 px','Pass'),('320','0 px','0','0 px','Pass')],[1.25,1.5,1.35,1.55,1.4])
para('The layout is designed defensively for unknown widths using fluid sizing, wrapping, container queries and standard input fallbacks. No finite test suite can guarantee identical rendering on every existing or future browser/device, so future content/layout changes should re-run the representative QA matrix rather than assuming previous measurements still apply.')

doc.add_page_break()
heading('12. Future-change checklist',1)
for x in [
'Change copy in index.html first; avoid hard-coding presentation logic into JavaScript.',
'Keep data-* hooks stable or update the matching JavaScript selector in the same commit.',
'When adding tables, populate data-label on every cell so narrow record-card mode remains understandable.',
'When adding AI/matrix nodes, give every node explicit numeric risk and oversight values.',
'Keep animation on transform/opacity where possible; preserve prefers-reduced-motion.',
'Do not add autoplay to governance or finite decks unless there is a compelling accessibility-safe reason.',
'Do not turn public risk signals into unsupported claims about a specific company.',
'Run check/build/postbuild and the viewport regression after layout/interaction changes.',
'If the production hostname changes, update canonical URL, Open Graph URL, robots.txt and sitemap.xml together.',
'Keep this DOCX and V17_IMPLEMENTATION_GUIDE.md updated if line numbers or major selectors change.'
]: bullet(x)

heading('Source notes',2)
para('External implementation references used for this release: CISA Known Exploited Vulnerabilities Catalog; CISA official cisagov/kev-data repository; Cloudflare Pages build configuration / static HTML deployment documentation; Vercel project configuration documentation. The production live section links directly to the authoritative CISA catalog.')

# Improve table row splitting behavior
for table in doc.tables:
    for row in table.rows:
        trPr=row._tr.get_or_add_trPr(); cant=OxmlElement('w:cantSplit'); trPr.append(cant)

# Save
doc.save(out)
# Make career-assets copy current.
(ROOT/'career-assets/Website_Code_Guide.docx').write_bytes(out.read_bytes())
print(out)
